from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import StreamingResponse
from datetime import datetime, timezone
import io

from db import db
from auth import get_current_user
from helpers import build_data_filter

router = APIRouter()


def _add_table_to_doc(doc, headers, rows):
    from docx.shared import Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = 1
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '2d2d3f',
            qn('w:val'): 'clear'
        })
        shading.append(shd)

    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8)
        if i % 2 == 1:
            for j in range(len(headers)):
                shading = table.rows[i + 1].cells[j]._element.get_or_add_tcPr()
                shd = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'f0f0f5',
                    qn('w:val'): 'clear'
                })
                shading.append(shd)


@router.get("/export/cars")
async def export_cars_word(current_user: dict = Depends(get_current_user)):
    from docx import Document
    from docx.shared import Cm
    from docx.enum.section import WD_ORIENT

    cars = await db.cars.find(build_data_filter(current_user, include_deleted=False), {"_id": 0}).to_list(5000)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    doc.add_heading('Arac Listesi', level=1)
    doc.add_paragraph(f'Toplam: {len(cars)} arac  |  Tarih: {datetime.now(timezone.utc).strftime("%d.%m.%Y")}')

    headers = ["Plaka", "Marka", "Model", "Yil", "KM", "Yakit", "Vites", "Durum", "Alis (TL)", "Satis (TL)", "Giris Tarihi"]

    rows = []
    for car in cars:
        rows.append([
            car.get("plate", "").upper(),
            car.get("brand", ""),
            car.get("model", ""),
            str(car.get("year", "")),
            f'{int(car.get("km", "0") or 0):,}'.replace(",", "."),
            car.get("fuel_type", ""),
            car.get("gear", ""),
            car.get("status", ""),
            f'{car.get("purchase_price", 0):,.0f}'.replace(",", "."),
            f'{car.get("sale_price", 0):,.0f}'.replace(",", "."),
            car.get("entry_date", ""),
        ])

    _add_table_to_doc(doc, headers, rows)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=araclar.docx"}
    )


@router.get("/export/customers")
async def export_customers_word(current_user: dict = Depends(get_current_user)):
    from docx import Document
    from docx.shared import Cm

    customers = await db.customers.find(build_data_filter(current_user, include_deleted=False), {"_id": 0}).to_list(5000)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    doc.add_heading('Musteri Listesi', level=1)
    doc.add_paragraph(f'Toplam: {len(customers)} musteri  |  Tarih: {datetime.now(timezone.utc).strftime("%d.%m.%Y")}')

    headers = ["Ad Soyad", "Telefon", "Tur", "Notlar", "Kayit Tarihi"]

    rows = []
    for c in customers:
        rows.append([
            c.get("name", ""),
            c.get("phone", ""),
            c.get("type", ""),
            c.get("notes", "")[:50],
            c.get("created_at", "")[:10],
        ])

    _add_table_to_doc(doc, headers, rows)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=musteriler.docx"}
    )


@router.get("/export/transactions")
async def export_transactions_word(current_user: dict = Depends(get_current_user)):
    from docx import Document
    from docx.shared import Cm
    from docx.enum.section import WD_ORIENT

    transactions_list = await db.transactions.find(build_data_filter(current_user, include_deleted=False), {"_id": 0}).to_list(5000)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    doc.add_heading('Islem Gecmisi', level=1)

    total_income = sum(t.get("amount", 0) for t in transactions_list if t.get("type") == "income")
    total_expense = sum(t.get("amount", 0) for t in transactions_list if t.get("type") == "expense")
    doc.add_paragraph(
        f'Toplam: {len(transactions_list)} islem  |  '
        f'Gelir: {total_income:,.0f} TL  |  Gider: {total_expense:,.0f} TL  |  '
        f'Net: {total_income - total_expense:,.0f} TL  |  '
        f'Tarih: {datetime.now(timezone.utc).strftime("%d.%m.%Y")}'
    )

    headers = ["Tarih", "Tur", "Kategori", "Aciklama", "Tutar (TL)"]

    rows = []
    for t in transactions_list:
        rows.append([
            t.get("date", ""),
            "Gelir" if t.get("type") == "income" else "Gider",
            t.get("category", ""),
            t.get("description", "")[:60],
            f'{t.get("amount", 0):,.0f}'.replace(",", "."),
        ])

    _add_table_to_doc(doc, headers, rows)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=islemler.docx"}
    )


@router.get("/export/expertise/{car_id}")
async def export_expertise_pdf(car_id: str, current_user: dict = Depends(get_current_user)):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm

    car = await db.cars.find_one({"id": car_id, "org_id": current_user.get("org_id", current_user["user_id"])}, {"_id": 0})
    if not car:
        raise HTTPException(status_code=404, detail="Car not found")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1.5*cm, bottomMargin=1.5*cm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title2', parent=styles['Title'], fontSize=18, spaceAfter=20)
    subtitle_style = ParagraphStyle('Sub', parent=styles['Normal'], fontSize=12, spaceAfter=10, textColor=colors.grey)

    elements = []

    elements.append(Paragraph("EKSPERTIZ RAPORU", title_style))
    elements.append(Paragraph(f"{car.get('brand', '')} {car.get('model', '')} - {car.get('year', '')} | {car.get('plate', '')}", subtitle_style))
    elements.append(Spacer(1, 15))

    info_data = [
        ["Marka", car.get("brand", ""), "Model", car.get("model", "")],
        ["Yil", str(car.get("year", "")), "Plaka", car.get("plate", "")],
        ["KM", car.get("km", ""), "Yakit", car.get("fuel_type", "")],
        ["Vites", car.get("gear", ""), "Motor", car.get("engine_type", "")],
        ["Il", car.get("province", ""), "Ilce", car.get("district", "")],
    ]

    info_table = Table(info_data, colWidths=[3*cm, 5.5*cm, 3*cm, 5.5*cm])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.Color(0.1, 0.1, 0.18)),
        ('BACKGROUND', (2, 0), (2, -1), colors.Color(0.1, 0.1, 0.18)),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.white),
        ('TEXTCOLOR', (2, 0), (2, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 20))

    elements.append(Paragraph("KAPORTA DURUMU", styles['Heading2']))

    status_labels = {"orijinal": "Orijinal", "boyali": "Boyali", "degisen": "Degisen", "lokal": "Lokal"}
    parts_data = [["Parca", "Durum"]]

    part_names = {
        "on_tampon": "On Tampon", "kaput": "Kaput", "sol_on_camurluk": "Sol On Camurluk",
        "sag_on_camurluk": "Sag On Camurluk", "sol_on_kapi": "Sol On Kapi",
        "sag_on_kapi": "Sag On Kapi", "tavan": "Tavan", "sol_arka_kapi": "Sol Arka Kapi",
        "sag_arka_kapi": "Sag Arka Kapi", "sol_arka_camurluk": "Sol Arka Camurluk",
        "sag_arka_camurluk": "Sag Arka Camurluk", "bagaj": "Bagaj", "arka_tampon": "Arka Tampon"
    }

    expertise_parts = car.get("expertise", {}).get("parts", {})
    for pid, pname in part_names.items():
        status = expertise_parts.get(pid, "orijinal")
        parts_data.append([pname, status_labels.get(status, status)])

    status_colors = {
        "orijinal": colors.Color(0.13, 0.77, 0.37),
        "boyali": colors.Color(0.92, 0.7, 0.05),
        "degisen": colors.Color(0.94, 0.27, 0.27),
        "lokal": colors.Color(0.23, 0.51, 0.94)
    }

    parts_table = Table(parts_data, colWidths=[8.5*cm, 8.5*cm])
    table_style_commands = [
        ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.1, 0.1, 0.18)),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
    ]

    for i, (pid, _) in enumerate(part_names.items()):
        status = expertise_parts.get(pid, "orijinal")
        color = status_colors.get(status, colors.white)
        table_style_commands.append(('BACKGROUND', (1, i+1), (1, i+1), color))

    parts_table.setStyle(TableStyle(table_style_commands))
    elements.append(parts_table)
    elements.append(Spacer(1, 20))

    elements.append(Paragraph("MEKANIK DURUM", styles['Heading2']))
    mechanical = car.get("expertise", {}).get("mechanical", {})
    mech_data = [
        ["Motor Durumu", mechanical.get("motor", "Orijinal")],
        ["Sanziman Durumu", mechanical.get("sanziman", "Orijinal")],
        ["Yuruyen Durumu", mechanical.get("yuruyen", "Orijinal")],
    ]
    mech_table = Table(mech_data, colWidths=[8.5*cm, 8.5*cm])
    mech_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (0, -1), colors.Color(0.1, 0.1, 0.18)),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.white),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(mech_table)
    elements.append(Spacer(1, 20))

    score_data = [
        ["Ekspertiz Puani (%)", str(car.get("expertise_score", 95))],
        ["Tramer Kayit Tutari (TL)", f"{car.get('tramer_amount', 0):,.0f}"],
    ]
    score_table = Table(score_data, colWidths=[8.5*cm, 8.5*cm])
    score_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (0, -1), colors.Color(0.1, 0.1, 0.18)),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.white),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(score_table)

    notes = car.get("expertise_notes", "")
    if notes:
        elements.append(Spacer(1, 15))
        elements.append(Paragraph("EKSPERTIZ NOTLARI", styles['Heading2']))
        elements.append(Paragraph(notes, styles['Normal']))

    elements.append(Spacer(1, 30))
    elements.append(Paragraph(f"Rapor Tarihi: {datetime.now(timezone.utc).strftime('%d.%m.%Y')}", subtitle_style))

    doc.build(elements)
    buffer.seek(0)

    filename = f"ekspertiz_{car.get('plate', 'rapor').replace(' ', '_')}.pdf"
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
