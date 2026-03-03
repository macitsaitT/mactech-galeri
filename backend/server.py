from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File, Query, Header
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.responses import StreamingResponse, Response
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional
import uuid
from datetime import datetime, timezone
import jwt
import bcrypt
import secrets
import hashlib
import requests
import io
from cryptography.fernet import Fernet
import base64

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# JWT Secret
JWT_SECRET = os.environ.get('JWT_SECRET', 'aslanbasoto-secret-key-2024')
JWT_ALGORITHM = "HS256"

# ==================== OBJECT STORAGE ====================
STORAGE_URL = "https://integrations.emergentagent.com/objstore/api/v1/storage"
EMERGENT_KEY = os.environ.get("EMERGENT_LLM_KEY")
APP_NAME = "aslanbasoto"
storage_key = None

def init_storage():
    global storage_key
    if storage_key:
        return storage_key
    resp = requests.post(f"{STORAGE_URL}/init", json={"emergent_key": EMERGENT_KEY}, timeout=30)
    resp.raise_for_status()
    storage_key = resp.json()["storage_key"]
    return storage_key

def put_object(path: str, data: bytes, content_type: str) -> dict:
    key = init_storage()
    resp = requests.put(
        f"{STORAGE_URL}/objects/{path}",
        headers={"X-Storage-Key": key, "Content-Type": content_type},
        data=data, timeout=120
    )
    resp.raise_for_status()
    return resp.json()

def get_object(path: str):
    key = init_storage()
    resp = requests.get(
        f"{STORAGE_URL}/objects/{path}",
        headers={"X-Storage-Key": key}, timeout=60
    )
    resp.raise_for_status()
    return resp.content, resp.headers.get("Content-Type", "application/octet-stream")

# ==================== ENCRYPTION ====================
def get_fernet():
    enc_key = os.environ.get('ENCRYPTION_KEY')
    if not enc_key:
        enc_key = Fernet.generate_key().decode()
        logger.info(f"Generated new encryption key. Store this: {enc_key}")
    if len(enc_key) == 44 and enc_key.endswith('='):
        return Fernet(enc_key.encode())
    key_bytes = hashlib.sha256(enc_key.encode()).digest()
    return Fernet(base64.urlsafe_b64encode(key_bytes))

fernet = None

def encrypt_value(value: str) -> str:
    global fernet
    if not fernet:
        fernet = get_fernet()
    if not value:
        return value
    return fernet.encrypt(value.encode()).decode()

def decrypt_value(value: str) -> str:
    global fernet
    if not fernet:
        fernet = get_fernet()
    if not value:
        return value
    try:
        return fernet.decrypt(value.encode()).decode()
    except Exception:
        return value

# Create the main app
app = FastAPI(title="Aslanbaş Oto CRM API")
api_router = APIRouter(prefix="/api")
security = HTTPBearer()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ==================== MODELS ====================

class UserCreate(BaseModel):
    email: str
    password: str
    company_name: str = "Aslanbaş Oto"
    phone: str = ""
    email_verified: bool = False
    role: str = "satis"  # admin, muhasebe, satis

class UserLogin(BaseModel):
    email: str
    password: str

class UserProfile(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: str
    company_name: str = "Aslanbaş Oto"
    phone: str = ""
    address: str = ""
    logo_url: str = ""
    theme: str = "dark"
    role: str = "admin"
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ProfileUpdate(BaseModel):
    company_name: Optional[str] = None
    phone: Optional[str] = None
    address: Optional[str] = None
    logo_url: Optional[str] = None
    theme: Optional[str] = None
    password: Optional[str] = None
    role: Optional[str] = None

class CarBase(BaseModel):
    brand: str
    model: str
    year: int
    plate: str
    km: str = ""
    vehicle_type: str = "Sedan"
    purchase_price: float = 0
    sale_price: float = 0
    description: str = ""
    status: str = "Stokta"
    entry_date: str = ""
    inspection_date: str = ""
    fuel_type: str = "Dizel"
    gear: str = "Otomatik"
    ownership: str = "stock"
    owner_name: str = ""
    owner_phone: str = ""
    commission_rate: float = 0
    photos: List[str] = []
    expertise: dict = {}
    package_info: str = ""
    engine_type: str = ""
    deposit_amount: float = 0
    sold_date: Optional[str] = None
    customer_id: Optional[str] = None
    customer_name: Optional[str] = None
    employee_share: float = 0
    insurance_start: str = ""
    insurance_end: str = ""
    province: str = ""
    district: str = ""
    expertise_score: int = 95
    tramer_amount: float = 0
    expertise_notes: str = ""

class CarCreate(CarBase):
    pass

class Car(CarBase):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    deleted: bool = False
    deleted_at: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class CustomerBase(BaseModel):
    name: str
    phone: str = ""
    type: str = "Potansiyel"
    notes: str = ""
    interested_car_id: str = ""

class CustomerCreate(CustomerBase):
    pass

class Customer(CustomerBase):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    deleted: bool = False
    deleted_at: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class TransactionBase(BaseModel):
    type: str  # income or expense
    category: str
    description: str = ""
    amount: float
    date: str
    car_id: Optional[str] = None
    employee_name: Optional[str] = None

class TransactionCreate(TransactionBase):
    pass

class Transaction(TransactionBase):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    deleted: bool = False
    deleted_at: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

# ==================== AUTH HELPERS ====================

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode(), hashed.encode())

def create_token(user_id: str, email: str, org_id: str = "", role: str = "admin") -> str:
    payload = {
        "user_id": user_id,
        "email": email,
        "org_id": org_id,
        "role": role,
        "exp": datetime.now(timezone.utc).timestamp() + 86400 * 30  # 30 days
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        token = credentials.credentials
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user_id = payload.get("user_id")
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token")
        return {
            "user_id": user_id,
            "email": payload.get("email"),
            "org_id": payload.get("org_id", user_id),
            "role": payload.get("role", "admin")
        }
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

# ==================== AUTH ROUTES ====================

@api_router.post("/auth/register")
async def register(user: UserCreate):
    existing = await db.users.find_one({"email": user.email})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Generate email verification code
    verification_code = str(secrets.randbelow(900000) + 100000)  # 6-digit code
    
    user_id = str(uuid.uuid4())
    # New registrations are always admin with their own org_id
    org_id = user_id
    user_doc = {
        "id": user_id,
        "email": user.email,
        "password_hash": hash_password(user.password),
        "company_name": user.company_name,
        "phone": user.phone,
        "address": "",
        "logo_url": "",
        "theme": "dark",
        "role": "admin",
        "org_id": org_id,
        "email_verified": False,
        "verification_code": verification_code,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.users.insert_one(user_doc)
    
    token = create_token(user_id, user.email, org_id, "admin")
    return {
        "token": token,
        "user": {"id": user_id, "email": user.email, "company_name": user.company_name, "phone": user.phone, "logo_url": "", "address": "", "role": "admin", "org_id": org_id},
        "verification_code": verification_code,
        "requires_verification": True
    }

@api_router.post("/auth/verify-email")
async def verify_email(data: dict):
    code = data.get("code", "")
    email = data.get("email", "")
    
    user = await db.users.find_one({"email": email})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    if user.get("email_verified"):
        return {"verified": True, "message": "Email already verified"}
    
    if user.get("verification_code") != code:
        raise HTTPException(status_code=400, detail="Invalid verification code")
    
    await db.users.update_one(
        {"email": email},
        {"$set": {"email_verified": True}, "$unset": {"verification_code": ""}}
    )
    
    return {"verified": True, "message": "Email verified successfully"}

@api_router.post("/auth/resend-verification")
async def resend_verification(data: dict):
    email = data.get("email", "")
    
    user = await db.users.find_one({"email": email})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    if user.get("email_verified"):
        return {"message": "Email already verified"}
    
    new_code = str(secrets.randbelow(900000) + 100000)
    await db.users.update_one({"email": email}, {"$set": {"verification_code": new_code}})
    
    return {"verification_code": new_code, "message": "New code generated"}

@api_router.post("/auth/login")
async def login(credentials: UserLogin):
    user = await db.users.find_one({"email": credentials.email})
    if not user:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    if not verify_password(credentials.password, user.get("password_hash", "")):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    org_id = user.get("org_id", user["id"])
    role = user.get("role", "admin")
    token = create_token(user["id"], user["email"], org_id, role)
    return {
        "token": token,
        "user": {
            "id": user["id"],
            "email": user["email"],
            "company_name": user.get("company_name", ""),
            "phone": user.get("phone", ""),
            "address": user.get("address", ""),
            "logo_url": user.get("logo_url", ""),
            "theme": user.get("theme", "dark"),
            "role": role,
            "org_id": org_id
        }
    }

@api_router.get("/auth/me")
async def get_me(current_user: dict = Depends(get_current_user)):
    user = await db.users.find_one({"id": current_user["user_id"]}, {"_id": 0, "password_hash": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return user

@api_router.put("/auth/profile")
async def update_profile(profile: ProfileUpdate, current_user: dict = Depends(get_current_user)):
    update_data = {k: v for k, v in profile.model_dump().items() if v is not None}
    
    if "password" in update_data:
        update_data["password_hash"] = hash_password(update_data.pop("password"))
    
    if update_data:
        update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
        await db.users.update_one({"id": current_user["user_id"]}, {"$set": update_data})
    
    user = await db.users.find_one({"id": current_user["user_id"]}, {"_id": 0, "password_hash": 0})
    return user

@api_router.delete("/auth/delete-account")
async def delete_account(current_user: dict = Depends(get_current_user)):
    user_id = current_user["user_id"]
    
    # Delete all user data
    await db.cars.delete_many({"user_id": user_id})
    await db.customers.delete_many({"user_id": user_id})
    await db.transactions.delete_many({"user_id": user_id})
    await db.users.delete_one({"id": user_id})
    
    return {"success": True, "message": "Account and all data deleted"}

# ==================== DATA QUERY HELPERS ====================

def build_data_filter(current_user: dict, extra_filter: dict = None, include_deleted: bool = True) -> dict:
    """Build MongoDB filter based on user role and org.
    - Admin: sees all org data
    - Muhasebe: sees all org data
    - Satis: sees only their own data
    """
    role = current_user.get("role", "admin")
    org_id = current_user.get("org_id", current_user["user_id"])
    
    if role == "satis":
        query = {"org_id": org_id, "created_by": current_user["user_id"]}
    else:
        query = {"org_id": org_id}
    
    if not include_deleted:
        query["deleted"] = False
    
    if extra_filter:
        query.update(extra_filter)
    return query

# ==================== CAR ROUTES ====================

@api_router.get("/cars", response_model=List[Car])
async def get_cars(created_by: str = None, current_user: dict = Depends(get_current_user)):
    extra = {}
    if created_by:
        extra["created_by"] = created_by
    query = build_data_filter(current_user, extra)
    cars = await db.cars.find(query, {"_id": 0}).to_list(1000)
    return cars

@api_router.post("/cars", response_model=Car)
async def create_car(car: CarCreate, current_user: dict = Depends(get_current_user)):
    car_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    car_doc = car.model_dump()
    car_doc.update({
        "id": car_id,
        "user_id": current_user["user_id"],
        "org_id": current_user.get("org_id", current_user["user_id"]),
        "created_by": current_user["user_id"],
        "deleted": False,
        "deleted_at": None,
        "created_at": now,
        "updated_at": now
    })
    
    await db.cars.insert_one(car_doc)
    
    # Return without _id
    result = await db.cars.find_one({"id": car_id}, {"_id": 0})
    return result

@api_router.put("/cars/{car_id}", response_model=Car)
async def update_car(car_id: str, car: CarCreate, current_user: dict = Depends(get_current_user)):
    org_id = current_user.get("org_id", current_user["user_id"])
    existing = await db.cars.find_one({"id": car_id, "org_id": org_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Car not found")
    
    update_data = car.model_dump()
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    await db.cars.update_one({"id": car_id}, {"$set": update_data})
    result = await db.cars.find_one({"id": car_id}, {"_id": 0})
    return result

@api_router.patch("/cars/{car_id}")
async def patch_car(car_id: str, updates: dict, current_user: dict = Depends(get_current_user)):
    org_id = current_user.get("org_id", current_user["user_id"])
    existing = await db.cars.find_one({"id": car_id, "org_id": org_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Car not found")
    
    updates["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.cars.update_one({"id": car_id}, {"$set": updates})
    result = await db.cars.find_one({"id": car_id}, {"_id": 0})
    return result

@api_router.delete("/cars/{car_id}")
async def delete_car(car_id: str, permanent: bool = False, current_user: dict = Depends(get_current_user)):
    org_id = current_user.get("org_id", current_user["user_id"])
    existing = await db.cars.find_one({"id": car_id, "org_id": org_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Car not found")
    
    if permanent:
        await db.cars.delete_one({"id": car_id})
        await db.transactions.delete_many({"car_id": car_id, "org_id": org_id})
    else:
        await db.cars.update_one({"id": car_id}, {"$set": {
            "deleted": True,
            "deleted_at": datetime.now(timezone.utc).isoformat()
        }})
        await db.transactions.update_many(
            {"car_id": car_id, "org_id": org_id},
            {"$set": {"deleted": True, "deleted_at": datetime.now(timezone.utc).isoformat()}}
        )
    
    return {"success": True}

@api_router.post("/cars/{car_id}/restore")
async def restore_car(car_id: str, current_user: dict = Depends(get_current_user)):
    org_id = current_user.get("org_id", current_user["user_id"])
    existing = await db.cars.find_one({"id": car_id, "org_id": org_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Car not found")
    
    await db.cars.update_one({"id": car_id}, {"$set": {"deleted": False, "deleted_at": None}})
    await db.transactions.update_many(
        {"car_id": car_id, "org_id": org_id},
        {"$set": {"deleted": False, "deleted_at": None}}
    )
    
    result = await db.cars.find_one({"id": car_id}, {"_id": 0})
    return result

# ==================== CUSTOMER ROUTES ====================

@api_router.get("/customers", response_model=List[Customer])
async def get_customers(created_by: str = None, current_user: dict = Depends(get_current_user)):
    extra = {}
    if created_by:
        extra["created_by"] = created_by
    query = build_data_filter(current_user, extra)
    customers = await db.customers.find(query, {"_id": 0}).to_list(1000)
    return customers

@api_router.post("/customers", response_model=Customer)
async def create_customer(customer: CustomerCreate, current_user: dict = Depends(get_current_user)):
    customer_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    customer_doc = customer.model_dump()
    customer_doc.update({
        "id": customer_id,
        "user_id": current_user["user_id"],
        "org_id": current_user.get("org_id", current_user["user_id"]),
        "created_by": current_user["user_id"],
        "deleted": False,
        "deleted_at": None,
        "created_at": now
    })
    
    await db.customers.insert_one(customer_doc)
    result = await db.customers.find_one({"id": customer_id}, {"_id": 0})
    return result

@api_router.put("/customers/{customer_id}", response_model=Customer)
async def update_customer(customer_id: str, customer: CustomerCreate, current_user: dict = Depends(get_current_user)):
    org_id = current_user.get("org_id", current_user["user_id"])
    existing = await db.customers.find_one({"id": customer_id, "org_id": org_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Customer not found")
    
    update_data = customer.model_dump()
    await db.customers.update_one({"id": customer_id}, {"$set": update_data})
    result = await db.customers.find_one({"id": customer_id}, {"_id": 0})
    return result

@api_router.delete("/customers/{customer_id}")
async def delete_customer(customer_id: str, permanent: bool = False, current_user: dict = Depends(get_current_user)):
    org_id = current_user.get("org_id", current_user["user_id"])
    existing = await db.customers.find_one({"id": customer_id, "org_id": org_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Customer not found")
    
    if permanent:
        await db.customers.delete_one({"id": customer_id})
    else:
        await db.customers.update_one({"id": customer_id}, {"$set": {
            "deleted": True,
            "deleted_at": datetime.now(timezone.utc).isoformat()
        }})
    
    return {"success": True}

@api_router.post("/customers/{customer_id}/restore")
async def restore_customer(customer_id: str, current_user: dict = Depends(get_current_user)):
    org_id = current_user.get("org_id", current_user["user_id"])
    existing = await db.customers.find_one({"id": customer_id, "org_id": org_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Customer not found")
    
    await db.customers.update_one({"id": customer_id}, {"$set": {"deleted": False, "deleted_at": None}})
    result = await db.customers.find_one({"id": customer_id}, {"_id": 0})
    return result

# ==================== TRANSACTION ROUTES ====================

@api_router.get("/transactions", response_model=List[Transaction])
async def get_transactions(created_by: str = None, current_user: dict = Depends(get_current_user)):
    extra = {}
    if created_by:
        extra["created_by"] = created_by
    query = build_data_filter(current_user, extra)
    transactions = await db.transactions.find(query, {"_id": 0}).to_list(5000)
    return transactions

@api_router.post("/transactions", response_model=Transaction)
async def create_transaction(transaction: TransactionCreate, current_user: dict = Depends(get_current_user)):
    transaction_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    transaction_doc = transaction.model_dump()
    transaction_doc.update({
        "id": transaction_id,
        "user_id": current_user["user_id"],
        "org_id": current_user.get("org_id", current_user["user_id"]),
        "created_by": current_user["user_id"],
        "deleted": False,
        "deleted_at": None,
        "created_at": now
    })
    
    await db.transactions.insert_one(transaction_doc)
    result = await db.transactions.find_one({"id": transaction_id}, {"_id": 0})
    return result

@api_router.put("/transactions/{transaction_id}")
async def update_transaction(transaction_id: str, updates: dict, current_user: dict = Depends(get_current_user)):
    org_id = current_user.get("org_id", current_user["user_id"])
    existing = await db.transactions.find_one({"id": transaction_id, "org_id": org_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Transaction not found")
    
    await db.transactions.update_one({"id": transaction_id}, {"$set": updates})
    result = await db.transactions.find_one({"id": transaction_id}, {"_id": 0})
    return result

@api_router.delete("/transactions/{transaction_id}")
async def delete_transaction(transaction_id: str, permanent: bool = False, current_user: dict = Depends(get_current_user)):
    org_id = current_user.get("org_id", current_user["user_id"])
    existing = await db.transactions.find_one({"id": transaction_id, "org_id": org_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Transaction not found")
    
    if permanent:
        await db.transactions.delete_one({"id": transaction_id})
    else:
        await db.transactions.update_one({"id": transaction_id}, {"$set": {
            "deleted": True,
            "deleted_at": datetime.now(timezone.utc).isoformat()
        }})
    
    return {"success": True}

# ==================== DASHBOARD STATS ====================

@api_router.get("/stats")
async def get_stats(current_user: dict = Depends(get_current_user)):
    query = build_data_filter(current_user, include_deleted=False)
    
    cars = await db.cars.find(query, {"_id": 0}).to_list(1000)
    
    stock_cars = [c for c in cars if c.get("status") == "Stokta"]
    consignment_cars = [c for c in cars if c.get("ownership") == "consignment" and c.get("status") != "Satıldı"]
    sold_cars = [c for c in cars if c.get("status") == "Satıldı"]
    deposit_cars = [c for c in cars if c.get("status") == "Kapora Alındı"]
    
    transactions = await db.transactions.find(build_data_filter(current_user, include_deleted=False), {"_id": 0}).to_list(5000)
    
    total_income = sum(t.get("amount", 0) for t in transactions if t.get("type") == "income")
    total_expense = sum(t.get("amount", 0) for t in transactions if t.get("type") == "expense")
    
    stock_value = sum(c.get("purchase_price", 0) for c in stock_cars)
    
    cust_query = build_data_filter(current_user, include_deleted=False)
    customers = await db.customers.count_documents(cust_query)
    
    return {
        "total_cars": len(cars),
        "stock_cars": len(stock_cars),
        "consignment_cars": len(consignment_cars),
        "sold_cars": len(sold_cars),
        "deposit_cars": len(deposit_cars),
        "total_income": total_income,
        "total_expense": total_expense,
        "net_profit": total_income - total_expense,
        "stock_value": stock_value,
        "total_customers": customers
    }

# ==================== USER MANAGEMENT (Admin) ====================

@api_router.get("/users")
async def get_users(current_user: dict = Depends(get_current_user)):
    """Get users in same org - admin only. Others get just their own info."""
    org_id = current_user.get("org_id", current_user["user_id"])
    role = current_user.get("role", "admin")
    
    if role != "admin":
        me = await db.users.find_one({"id": current_user["user_id"]}, {"_id": 0, "password_hash": 0, "verification_code": 0})
        return [me] if me else []
    
    users = await db.users.find({"org_id": org_id}, {"_id": 0, "password_hash": 0, "verification_code": 0}).to_list(100)
    return users

@api_router.post("/users")
async def create_user(user: UserCreate, current_user: dict = Depends(get_current_user)):
    """Create a new user - admin only. Inherits admin's org_id."""
    if current_user.get("role", "admin") != "admin":
        raise HTTPException(status_code=403, detail="Sadece admin kullanıcı ekleyebilir")
    
    existing = await db.users.find_one({"email": user.email})
    if existing:
        raise HTTPException(status_code=400, detail="Bu email zaten kayıtlı")
    
    org_id = current_user.get("org_id", current_user["user_id"])
    user_id = str(uuid.uuid4())
    user_doc = {
        "id": user_id,
        "email": user.email,
        "password_hash": hash_password(user.password),
        "company_name": user.company_name,
        "phone": user.phone,
        "address": "",
        "logo_url": "",
        "theme": "dark",
        "role": user.role,
        "org_id": org_id,
        "email_verified": True,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.users.insert_one(user_doc)
    
    return {"id": user_id, "email": user.email, "company_name": user.company_name, "phone": user.phone, "role": user.role, "org_id": org_id}

@api_router.put("/users/{user_id}")
async def update_user(user_id: str, updates: dict, current_user: dict = Depends(get_current_user)):
    """Update user role/info - admin only, same org."""
    if current_user.get("role", "admin") != "admin":
        raise HTTPException(status_code=403, detail="Sadece admin kullanıcı düzenleyebilir")
    
    org_id = current_user.get("org_id", current_user["user_id"])
    target = await db.users.find_one({"id": user_id, "org_id": org_id})
    if not target:
        raise HTTPException(status_code=404, detail="Kullanıcı bulunamadı")
    
    allowed = {"role", "company_name", "phone", "email"}
    safe_updates = {k: v for k, v in updates.items() if k in allowed and v is not None}
    
    if "password" in updates and updates["password"]:
        safe_updates["password_hash"] = hash_password(updates["password"])
    
    if safe_updates:
        safe_updates["updated_at"] = datetime.now(timezone.utc).isoformat()
        await db.users.update_one({"id": user_id}, {"$set": safe_updates})
    
    updated = await db.users.find_one({"id": user_id}, {"_id": 0, "password_hash": 0})
    return updated

@api_router.delete("/users/{user_id}")
async def delete_user(user_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a user - admin only, same org. Cannot delete self."""
    if current_user.get("role", "admin") != "admin":
        raise HTTPException(status_code=403, detail="Sadece admin kullanıcı silebilir")
    
    if user_id == current_user["user_id"]:
        raise HTTPException(status_code=400, detail="Kendinizi silemezsiniz")
    
    org_id = current_user.get("org_id", current_user["user_id"])
    target = await db.users.find_one({"id": user_id, "org_id": org_id})
    if not target:
        raise HTTPException(status_code=404, detail="Kullanıcı bulunamadı")
    
    await db.users.delete_one({"id": user_id})
    return {"success": True}

@api_router.get("/employees")
async def get_employees(current_user: dict = Depends(get_current_user)):
    """Get all employees in same org for dropdown selections."""
    org_id = current_user.get("org_id", current_user["user_id"])
    users = await db.users.find({"org_id": org_id}, {"_id": 0, "password_hash": 0, "verification_code": 0}).to_list(100)
    return [{"id": u["id"], "email": u["email"], "name": u.get("company_name", u["email"]), "phone": u.get("phone", ""), "role": u.get("role", "satis")} for u in users]

@api_router.get("/org-users")
async def get_org_users(current_user: dict = Depends(get_current_user)):
    """Get all users in same org for muhasebe filtering."""
    org_id = current_user.get("org_id", current_user["user_id"])
    users = await db.users.find({"org_id": org_id}, {"_id": 0, "password_hash": 0, "verification_code": 0}).to_list(100)
    return [{"id": u["id"], "name": u.get("company_name", u["email"]), "role": u.get("role", "satis")} for u in users]

# ==================== HEALTH CHECK ====================

@api_router.get("/")
async def root():
    return {"message": "Aslanbaş Oto CRM API", "status": "running"}

@api_router.get("/health")
async def health():
    return {"status": "healthy"}

# ==================== FILE UPLOAD ====================

@api_router.post("/upload")
async def upload_file(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    ext = file.filename.split(".")[-1].lower() if "." in file.filename else "bin"
    allowed = {"jpg", "jpeg", "png", "gif", "webp"}
    if ext not in allowed:
        raise HTTPException(status_code=400, detail="Only image files allowed")
    
    data = await file.read()
    if len(data) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large (max 10MB)")
    
    path = f"{APP_NAME}/uploads/{current_user['user_id']}/{uuid.uuid4()}.{ext}"
    
    mime = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png", "gif": "image/gif", "webp": "image/webp"}
    result = put_object(path, data, mime.get(ext, "application/octet-stream"))
    
    file_doc = {
        "id": str(uuid.uuid4()),
        "storage_path": result["path"],
        "original_filename": file.filename,
        "content_type": file.content_type,
        "size": result.get("size", len(data)),
        "user_id": current_user["user_id"],
        "is_deleted": False,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.files.insert_one(file_doc)
    
    return {"id": file_doc["id"], "path": result["path"], "filename": file.filename}

@api_router.get("/files/{file_path:path}")
async def download_file(file_path: str, auth: str = Query(None), authorization: str = Header(None)):
    auth_header = authorization or (f"Bearer {auth}" if auth else None)
    if not auth_header:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    try:
        token = auth_header.replace("Bearer ", "")
        jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid token")
    
    record = await db.files.find_one({"storage_path": file_path, "is_deleted": False})
    if not record:
        raise HTTPException(status_code=404, detail="File not found")
    
    data, content_type = get_object(file_path)
    return Response(content=data, media_type=record.get("content_type", content_type))

# ==================== WORD EXPORT ====================

def _add_table_to_doc(doc, headers, rows):
    """Helper to add a styled table to a Word document."""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = 1  # CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        # Dark background
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '1a1a2e',
            qn('w:val'): 'clear'
        })
        shading.append(shd)
    
    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

@api_router.get("/export/cars")
async def export_cars_word(current_user: dict = Depends(get_current_user)):
    from docx import Document
    from docx.shared import Pt
    
    cars = await db.cars.find(build_data_filter(current_user, include_deleted=False), {"_id": 0}).to_list(5000)
    
    doc = Document()
    doc.add_heading('Araç Listesi', level=1)
    doc.add_paragraph(f'Toplam: {len(cars)} araç | Tarih: {datetime.now(timezone.utc).strftime("%d.%m.%Y")}')
    
    headers = ["Plaka", "Marka", "Model", "Yıl", "KM", "Yakıt", "Vites", "Durum",
               "Alış Fiyatı", "Satış Fiyatı", "Ekspertiz", "Tramer", "Giriş Tarihi"]
    
    rows = []
    for car in cars:
        rows.append([
            car.get("plate", ""),
            car.get("brand", ""),
            car.get("model", ""),
            str(car.get("year", "")),
            car.get("km", ""),
            car.get("fuel_type", ""),
            car.get("gear", ""),
            car.get("status", ""),
            f'{car.get("purchase_price", 0):,.0f} TL',
            f'{car.get("sale_price", 0):,.0f} TL',
            str(car.get("expertise_score", 0)),
            f'{car.get("tramer_amount", 0):,.0f} TL',
            car.get("entry_date", ""),
        ])
    
    _add_table_to_doc(doc, headers, rows)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=araclar.docx"}
    )

@api_router.get("/export/customers")
async def export_customers_word(current_user: dict = Depends(get_current_user)):
    from docx import Document
    
    customers = await db.customers.find(build_data_filter(current_user, include_deleted=False), {"_id": 0}).to_list(5000)
    
    doc = Document()
    doc.add_heading('Müşteri Listesi', level=1)
    doc.add_paragraph(f'Toplam: {len(customers)} müşteri | Tarih: {datetime.now(timezone.utc).strftime("%d.%m.%Y")}')
    
    headers = ["Ad Soyad", "Telefon", "Tür", "Notlar", "Kayıt Tarihi"]
    
    rows = []
    for c in customers:
        rows.append([
            c.get("name", ""),
            c.get("phone", ""),
            c.get("type", ""),
            c.get("notes", ""),
            c.get("created_at", "")[:10],
        ])
    
    _add_table_to_doc(doc, headers, rows)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=musteriler.docx"}
    )

@api_router.get("/export/transactions")
async def export_transactions_word(current_user: dict = Depends(get_current_user)):
    from docx import Document
    
    transactions_list = await db.transactions.find(build_data_filter(current_user, include_deleted=False), {"_id": 0}).to_list(5000)
    
    doc = Document()
    doc.add_heading('İşlem Geçmişi', level=1)
    doc.add_paragraph(f'Toplam: {len(transactions_list)} işlem | Tarih: {datetime.now(timezone.utc).strftime("%d.%m.%Y")}')
    
    headers = ["Tarih", "Tür", "Kategori", "Açıklama", "Tutar (TL)"]
    
    rows = []
    for t in transactions_list:
        rows.append([
            t.get("date", ""),
            "Gelir" if t.get("type") == "income" else "Gider",
            t.get("category", ""),
            t.get("description", ""),
            f'{t.get("amount", 0):,.0f}',
        ])
    
    _add_table_to_doc(doc, headers, rows)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=islemler.docx"}
    )

# ==================== PDF EXPERTISE REPORT ====================

@api_router.get("/export/expertise/{car_id}")
async def export_expertise_pdf(car_id: str, current_user: dict = Depends(get_current_user)):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    
    car = await db.cars.find_one({"id": car_id, "org_id": current_user.get("org_id", current_user["user_id"])}, {"_id": 0})
    if not car:
        raise HTTPException(status_code=404, detail="Car not found")
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1.5*cm, bottomMargin=1.5*cm)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title2', parent=styles['Title'], fontSize=18, spaceAfter=20)
    subtitle_style = ParagraphStyle('Sub', parent=styles['Normal'], fontSize=12, spaceAfter=10, textColor=colors.grey)
    
    elements = []
    
    # Title
    elements.append(Paragraph("EKSPERTIZ RAPORU", title_style))
    elements.append(Paragraph(f"{car.get('brand', '')} {car.get('model', '')} - {car.get('year', '')} | {car.get('plate', '')}", subtitle_style))
    elements.append(Spacer(1, 15))
    
    # Vehicle Info Table
    info_data = [
        ["Marka", car.get("brand", ""), "Model", car.get("model", "")],
        ["Yil", str(car.get("year", "")), "Plaka", car.get("plate", "")],
        ["KM", car.get("km", ""), "Yakit", car.get("fuel_type", "")],
        ["Vites", car.get("gear", ""), "Motor", car.get("engine_type", "")],
        ["Il", car.get("province", ""), "Ilce", car.get("district", "")],
    ]
    
    info_table = Table(info_data, colWidths=[3*cm, 5.5*cm, 3*cm, 5.5*cm])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.Color(0.1, 0.1, 0.18)),
        ('BACKGROUND', (2, 0), (2, -1), colors.Color(0.1, 0.1, 0.18)),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.white),
        ('TEXTCOLOR', (2, 0), (2, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 20))
    
    # Expertise Parts
    elements.append(Paragraph("KAPORTA DURUMU", styles['Heading2']))
    
    status_labels = {"orijinal": "Orijinal", "boyali": "Boyali", "degisen": "Degisen", "lokal": "Lokal"}
    parts_data = [["Parca", "Durum"]]
    
    part_names = {
        "on_tampon": "On Tampon", "kaput": "Kaput", "sol_on_camurluk": "Sol On Camurluk",
        "sag_on_camurluk": "Sag On Camurluk", "sol_on_kapi": "Sol On Kapi",
        "sag_on_kapi": "Sag On Kapi", "tavan": "Tavan", "sol_arka_kapi": "Sol Arka Kapi",
        "sag_arka_kapi": "Sag Arka Kapi", "sol_arka_camurluk": "Sol Arka Camurluk",
        "sag_arka_camurluk": "Sag Arka Camurluk", "bagaj": "Bagaj", "arka_tampon": "Arka Tampon"
    }
    
    expertise_parts = car.get("expertise", {}).get("parts", {})
    for pid, pname in part_names.items():
        status = expertise_parts.get(pid, "orijinal")
        parts_data.append([pname, status_labels.get(status, status)])
    
    status_colors = {
        "orijinal": colors.Color(0.13, 0.77, 0.37),
        "boyali": colors.Color(0.92, 0.7, 0.05),
        "degisen": colors.Color(0.94, 0.27, 0.27),
        "lokal": colors.Color(0.23, 0.51, 0.94)
    }
    
    parts_table = Table(parts_data, colWidths=[8.5*cm, 8.5*cm])
    table_style_commands = [
        ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.1, 0.1, 0.18)),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
    ]
    
    for i, (pid, _) in enumerate(part_names.items()):
        status = expertise_parts.get(pid, "orijinal")
        color = status_colors.get(status, colors.white)
        table_style_commands.append(('BACKGROUND', (1, i+1), (1, i+1), color))
    
    parts_table.setStyle(TableStyle(table_style_commands))
    elements.append(parts_table)
    elements.append(Spacer(1, 20))
    
    # Mechanical Status
    elements.append(Paragraph("MEKANIK DURUM", styles['Heading2']))
    mechanical = car.get("expertise", {}).get("mechanical", {})
    mech_data = [
        ["Motor Durumu", mechanical.get("motor", "Orijinal")],
        ["Sanziman Durumu", mechanical.get("sanziman", "Orijinal")],
        ["Yuruyen Durumu", mechanical.get("yuruyen", "Orijinal")],
    ]
    mech_table = Table(mech_data, colWidths=[8.5*cm, 8.5*cm])
    mech_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (0, -1), colors.Color(0.1, 0.1, 0.18)),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.white),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(mech_table)
    elements.append(Spacer(1, 20))
    
    # Score and Tramer
    score_data = [
        ["Ekspertiz Puani (%)", str(car.get("expertise_score", 95))],
        ["Tramer Kayit Tutari (TL)", f"{car.get('tramer_amount', 0):,.0f}"],
    ]
    score_table = Table(score_data, colWidths=[8.5*cm, 8.5*cm])
    score_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (0, -1), colors.Color(0.1, 0.1, 0.18)),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.white),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(score_table)
    
    # Notes
    notes = car.get("expertise_notes", "")
    if notes:
        elements.append(Spacer(1, 15))
        elements.append(Paragraph("EKSPERTIZ NOTLARI", styles['Heading2']))
        elements.append(Paragraph(notes, styles['Normal']))
    
    elements.append(Spacer(1, 30))
    elements.append(Paragraph(f"Rapor Tarihi: {datetime.now(timezone.utc).strftime('%d.%m.%Y')}", subtitle_style))
    
    doc.build(elements)
    buffer.seek(0)
    
    filename = f"ekspertiz_{car.get('plate', 'rapor').replace(' ', '_')}.pdf"
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

# ==================== ENCRYPTION ENDPOINTS ====================

@api_router.post("/encrypt-customer/{customer_id}")
async def encrypt_customer_data(customer_id: str, current_user: dict = Depends(get_current_user)):
    customer = await db.customers.find_one({"id": customer_id, "org_id": current_user.get("org_id", current_user["user_id"])})
    if not customer:
        raise HTTPException(status_code=404, detail="Customer not found")
    
    encrypted_fields = {}
    if customer.get("phone"):
        encrypted_fields["phone_encrypted"] = encrypt_value(customer["phone"])
    if customer.get("notes"):
        encrypted_fields["notes_encrypted"] = encrypt_value(customer["notes"])
    
    encrypted_fields["is_encrypted"] = True
    await db.customers.update_one({"id": customer_id}, {"$set": encrypted_fields})
    
    return {"success": True, "message": "Customer data encrypted"}

# ==================== APPOINTMENTS (TEST DRIVE) ====================

class AppointmentBase(BaseModel):
    model_config = ConfigDict(extra="ignore")
    title: str
    customer_name: str = ""
    customer_phone: str = ""
    car_id: str = ""
    car_info: str = ""
    date: str = ""
    time: str = ""
    notes: str = ""
    status: str = "Bekliyor"

@api_router.get("/appointments")
async def get_appointments(current_user: dict = Depends(get_current_user)):
    query = build_data_filter(current_user)
    query.pop("deleted", None)  # appointments don't have soft delete
    appointments = await db.appointments.find(query, {"_id": 0}).sort("date", 1).to_list(1000)
    return appointments

@api_router.post("/appointments")
async def create_appointment(appointment: AppointmentBase, current_user: dict = Depends(get_current_user)):
    doc = appointment.model_dump()
    doc["id"] = str(uuid.uuid4())
    doc["user_id"] = current_user["user_id"]
    doc["org_id"] = current_user.get("org_id", current_user["user_id"])
    doc["created_by"] = current_user["user_id"]
    doc["created_at"] = datetime.now(timezone.utc).isoformat()
    await db.appointments.insert_one(doc)
    doc.pop("_id", None)
    return doc

@api_router.put("/appointments/{appointment_id}")
async def update_appointment(appointment_id: str, data: dict, current_user: dict = Depends(get_current_user)):
    org_id = current_user.get("org_id", current_user["user_id"])
    data.pop("_id", None)
    data.pop("id", None)
    data["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.appointments.update_one(
        {"id": appointment_id, "org_id": org_id},
        {"$set": data}
    )
    updated = await db.appointments.find_one({"id": appointment_id}, {"_id": 0})
    return updated

@api_router.delete("/appointments/{appointment_id}")
async def delete_appointment(appointment_id: str, current_user: dict = Depends(get_current_user)):
    org_id = current_user.get("org_id", current_user["user_id"])
    await db.appointments.delete_one({"id": appointment_id, "org_id": org_id})
    return {"success": True}

# Include router
app.include_router(api_router)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("startup")
async def startup():
    try:
        init_storage()
        logger.info("Object Storage initialized")
    except Exception as e:
        logger.warning(f"Storage init deferred: {e}")
    
    # Migration: set role=admin and org_id=user_id for users without these fields
    await db.users.update_many(
        {"role": {"$exists": False}},
        {"$set": {"role": "admin"}}
    )
    
    # Set org_id for users without it (org_id = their own user_id for admins)
    users_without_org = await db.users.find({"org_id": {"$exists": False}}, {"_id": 0, "id": 1}).to_list(1000)
    for u in users_without_org:
        await db.users.update_one({"id": u["id"]}, {"$set": {"org_id": u["id"]}})
    
    # Set org_id and created_by on data documents that don't have them
    for collection_name in ["cars", "customers", "transactions", "appointments"]:
        col = db[collection_name]
        docs = await col.find({"org_id": {"$exists": False}, "user_id": {"$exists": True}}, {"_id": 0, "id": 1, "user_id": 1}).to_list(10000)
        for d in docs:
            await col.update_one(
                {"id": d["id"]},
                {"$set": {"org_id": d["user_id"], "created_by": d["user_id"]}}
            )
    
    logger.info("Data migration complete (org_id, created_by, role)")

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
